"""Document processing utilities for extracting text from various file formats."""

import logging
from typing import Optional, BinaryIO
from pathlib import Path
import io

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles text extraction from various document formats."""
    
    SUPPORTED_FORMATS = ['.txt', '.pdf', '.docx', '.doc']
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
    
    @staticmethod
    def get_supported_formats() -> list[str]:
        """
        Get list of supported document formats.
        
        Returns:
            List of supported file extensions
        """
        return DocumentProcessor.SUPPORTED_FORMATS.copy()
    
    @staticmethod
    def is_supported_format(filename: str) -> bool:
        """
        Check if a file format is supported.
        
        Args:
            filename: Name of the file
        
        Returns:
            True if format is supported, False otherwise
        """
        ext = Path(filename).suffix.lower()
        return ext in DocumentProcessor.SUPPORTED_FORMATS
    
    @staticmethod
    async def extract_text(
        file_content: bytes,
        filename: str,
        max_length: Optional[int] = None
    ) -> str:
        """
        Extract text from a document file.
        
        Args:
            file_content: Binary content of the file
            filename: Name of the file (used to determine format)
            max_length: Optional maximum text length to extract
        
        Returns:
            Extracted text content
        
        Raises:
            ValueError: If file format is not supported or file is too large
            Exception: If text extraction fails
        """
        try:
            # Validate file size
            if len(file_content) > DocumentProcessor.MAX_FILE_SIZE:
                raise ValueError(
                    f"File size exceeds maximum of {DocumentProcessor.MAX_FILE_SIZE / (1024*1024):.1f} MB"
                )
            
            # Get file extension
            ext = Path(filename).suffix.lower()
            
            if not DocumentProcessor.is_supported_format(filename):
                raise ValueError(
                    f"Unsupported file format: {ext}. Supported formats: {DocumentProcessor.SUPPORTED_FORMATS}"
                )
            
            logger.info(f"Extracting text from {filename} ({len(file_content)} bytes)")
            
            # Extract text based on format
            if ext == '.txt':
                text = DocumentProcessor._extract_text_from_txt(file_content)
            elif ext == '.pdf':
                text = DocumentProcessor._extract_text_from_pdf(file_content)
            elif ext in ['.docx', '.doc']:
                text = DocumentProcessor._extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported format: {ext}")
            
            # Validate extracted text
            if not text or not text.strip():
                raise ValueError("No text could be extracted from the document")
            
            # Truncate if needed
            if max_length and len(text) > max_length:
                logger.warning(f"Text truncated from {len(text)} to {max_length} characters")
                text = text[:max_length]
            
            logger.info(f"Successfully extracted {len(text)} characters from {filename}")
            return text.strip()
        
        except ValueError as e:
            logger.error(f"Validation error extracting text from {filename}: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise Exception(f"Failed to extract text from document: {str(e)}")
    
    @staticmethod
    def _extract_text_from_txt(file_content: bytes) -> str:
        """
        Extract text from a plain text file.
        
        Args:
            file_content: Binary content of the file
        
        Returns:
            Extracted text
        """
        try:
            # Try UTF-8 first
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1 if UTF-8 fails
                logger.warning("UTF-8 decoding failed, trying latin-1")
                return file_content.decode('latin-1')
        except Exception as e:
            raise Exception(f"Failed to decode text file: {str(e)}")
    
    @staticmethod
    def _extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_content: Binary content of the file
        
        Returns:
            Extracted text
        """
        try:
            from PyPDF2 import PdfReader
            
            # Create a file-like object from bytes
            pdf_file = io.BytesIO(file_content)
            
            # Read PDF
            reader = PdfReader(pdf_file)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
            
            if not text_parts:
                raise ValueError("No text could be extracted from PDF")
            
            return "\n\n".join(text_parts)
        
        except ImportError:
            raise Exception("PyPDF2 library is not installed. Install it with: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def _extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_content: Binary content of the file
        
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            # Create a file-like object from bytes
            docx_file = io.BytesIO(file_content)
            
            # Read DOCX
            doc = Document(docx_file)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            if not text_parts:
                raise ValueError("No text could be extracted from DOCX")
            
            return "\n\n".join(text_parts)
        
        except ImportError:
            raise Exception("python-docx library is not installed. Install it with: pip install python-docx")
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def split_text_into_chunks(
        text: str,
        max_chunk_size: int = 4000,
        overlap: int = 200
    ) -> list[str]:
        """
        Split text into chunks for processing long documents.
        
        Args:
            text: Text to split
            max_chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks
        
        Returns:
            List of text chunks
        """
        if len(text) <= max_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find the end of this chunk
            end = start + max_chunk_size
            
            # If this is not the last chunk, try to break at a sentence or word boundary
            if end < len(text):
                # Look for sentence boundary (. ! ?)
                sentence_end = max(
                    text.rfind('. ', start, end),
                    text.rfind('! ', start, end),
                    text.rfind('? ', start, end)
                )
                
                if sentence_end > start:
                    end = sentence_end + 1
                else:
                    # Look for word boundary
                    space_pos = text.rfind(' ', start, end)
                    if space_pos > start:
                        end = space_pos
            
            # Add chunk
            chunks.append(text[start:end].strip())
            
            # Move start position with overlap
            start = end - overlap if end < len(text) else end
        
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks
